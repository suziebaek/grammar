import streamlit as st
import random
from pathlib import Path
import google.generativeai as genai
from openai import OpenAI, AzureOpenAI
import pandas as pd
import re
import time
import io
from docx import Document # 🚀 [추가] 워드 다운로드를 위한 라이브러리 (pip install python-docx 필요)
from docx.shared import Pt  # 🚀 [추가] 폰트 크기(Pt) 조절을 위한 모듈
from validator import validate_question_llm, validate_batch_llm  # <--- 이 줄이 반드시 있어야 합니다!
from datetime import datetime
from prompts import build_generation_prompt as build_prompt_h
from prompts_e import build_generation_prompt_e
from prompts import build_retry_prompt
# 🚀 팝업창(모달)을 띄우기 위한 스트림릿 데코레이터 함수 추가
@st.dialog("⚠️ 필수 설정 누락")
def show_validator_warning():
    st.warning("LLM 검증기 작동 여부가 선택되지 않았습니다.\n\n버튼 옆의 메뉴를 **'❌ 검증기 끄기'** 또는 **'✅ 검증기 켜기'**로 변경한 후 다시 생성해 주세요.")
    
    if st.button("확인", type="primary", use_container_width=True):
        st.rerun()
        
TOPIC_LIST = [
    "역사", "우주 탐사", "심해 생물", "인공지능", "날씨", "일상 생활(여행)",
    "일상 생활(방과후 활동)", "일상 생활(규칙)", "스포츠 과학", "농업", "일상 생활(쇼핑)", "우정",
    "일상 생활(가족)", "기후 변화", "건축 공학", "야생 동물 보호",
    "지구 온난화", "일상 생활(운전)", "심리학", "학교 생활", "문학", "일상 생활(건강)",
    "로봇", "환경", "지리", "종교", "소셜 미디어", "요리", "일상 생활(식사)"
    
]

# 🚀 1. 난이도 분배 함수 (H레벨 최대 7점 제한 적용)
def get_difficulty_combs(is_e_level):
    ALL_COMBS = [(a, b, c) for a in (0, 1, 2, 3) for b in (0, 1, 2, 3) for c in (0, 1, 2, 3)]
    
    if is_e_level:
        # 🔵 E 레벨 (높은 레벨 / 심화)
        easy = [c for c in ALL_COMBS if sum(c) <= 3]       # 하: 0~3점
        mid  = [c for c in ALL_COMBS if 4 <= sum(c) <= 6]  # 중: 4~6점
        hard = [c for c in ALL_COMBS if 7 <= sum(c) <= 9]  # 상: 7~9점
    else:
        # 🟡 H 레벨 (낮은 레벨 / 기본 - 최대 7점까지만 허용)
        easy = [c for c in ALL_COMBS if sum(c) <= 2]       # 하: 0~2점
        mid  = [c for c in ALL_COMBS if 3 <= sum(c) <= 5]  # 중: 3~5점
        hard = [c for c in ALL_COMBS if 6 <= sum(c) <= 7]  # 상: 6~7점 (8, 9점 원천 차단)
        
    return easy, mid, hard



def add_paragraph_with_tags(doc_or_element, text):
    p = doc_or_element.add_paragraph()
    # <u>태그 단위로 텍스트 쪼개기
    parts = re.split(r'(<u>.*?</u>)', text)
    for part in parts:
        if part.startswith('<u>') and part.endswith('</u>'):
            run = p.add_run(part[3:-4]) # <u> </u> 떼어내고 알맹이만
            run.underline = True        # 워드 밑줄 속성 ON
        else:
            p.add_run(part)
            
# 🚀 [수정] 기존 create_word_document 함수 통째로 교체
def create_word_document(history_data, is_multiple=False):
    doc = Document()
    
    # 🚀 [추가] 챕터(생성정보) 글자 크기를 10pt로 강제하는 헬퍼 함수
    def add_custom_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10) # 글자 크기 10
        run.bold = True
        
    # 🚀 [추가] 워드에 쓸 때 [발문], [보기/지문] 등 브라켓을 싹 지우는 함수
    def clean_and_add_text(text):
        clean_text = text.replace("[발문]\n", "").replace("[보기/지문]\n", "").replace("[정답]\n", "").replace("[해설]\n", "")
        add_paragraph_with_tags(doc, clean_text.strip())

    if not is_multiple:
        entry = history_data
        add_custom_heading(f"생성 정보: {entry['major']} > {entry['mid']} > {entry['minor']}")
        doc.add_paragraph(f"난이도: {entry['difficulty']}")
        
        prev_type = None
# 기존: add_paragraph_with_tags(doc, r['text'])

# 🚀 단일 세트 처리 부분 (수정 후)
        for r in entry["results"]:
            if r['type'] != prev_type:
                doc.add_heading(f"🟦 {r['type']} 유형", level=1)
                prev_type = r['type']
                # UI용 text 대신 다운로드용 dl_text 삽입
            add_paragraph_with_tags(doc, r.get('dl_text', r['text']))


    else:
        doc.add_heading("전체 생성 문제 통합본", 0)
        for i, h in enumerate(history_data):
            add_custom_heading(f"[세트 {i+1}] {h['major']} > {h['mid']} > {h['minor']}")
            doc.add_paragraph(f"난이도: {h['difficulty']}")
            
            prev_type = None
# 🚀 다중 세트(전체) 처리 부분 (수정 후)
            for r in h["results"]:
                if r['type'] != prev_type:
                    doc.add_heading(f"🟦 {r['type']} 유형", level=2) 
                    prev_type = r['type']
                # UI용 text 대신 다운로드용 dl_text 삽입
                add_paragraph_with_tags(doc, r.get('dl_text', r['text']))
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ── 🎯 전역 설정: 구글 시트 탭별 GID URL 하드코딩 ────────────────────────
# [H레벨 시트 URL]
QUESTIONS_SHEET_URL = "https://docs.google.com/spreadsheets/d/1gSMH96-BB8sjs4FbNy8bb_KSnP8zOpBQPQ_6Q4ylZ90/edit?gid=939067680#gid=939067680"
CONCEPTS_SHEET_URL = "https://docs.google.com/spreadsheets/d/1gSMH96-BB8sjs4FbNy8bb_KSnP8zOpBQPQ_6Q4ylZ90/edit?gid=0#gid=0"

# 🚀 [추가] [E레벨 시트 URL] (선생님이 만드신 E레벨 전용 시트 URL로 교체하세요!)
E_QUESTIONS_SHEET_URL = "https://docs.google.com/spreadsheets/d/1gSMH96-BB8sjs4FbNy8bb_KSnP8zOpBQPQ_6Q4ylZ90/edit?gid=900494344#gid=900494344"
E_CONCEPTS_SHEET_URL = "https://docs.google.com/spreadsheets/d/1gSMH96-BB8sjs4FbNy8bb_KSnP8zOpBQPQ_6Q4ylZ90/edit?gid=1215243386#gid=1215243386"

# ── 페이지 설정 ───────────────────────────────────────────
st.set_page_config(
    page_title="영어 기출 문제 생성기",
    page_icon="📝",
    layout="wide",
)

BASE = Path(__file__).parent
# 선지 재정렬
def sort_options(passage_text, ans_text, exp_text):
    match = re.search(r'(.*?)(①.*)', passage_text, re.DOTALL)
    if not match: return passage_text, ans_text, exp_text
    
    passage, opts_raw = match.group(1), match.group(2)
    opts = re.split(r'([①-⑤])', opts_raw)[1:]
    
    if len(opts) != 10: return passage_text, ans_text, exp_text
    
    old_opt_dict = {opts[i]: opts[i+1].strip() for i in range(0, len(opts), 2)}
    ans_match = re.search(r'[①-⑤]', ans_text)
    old_ans = ans_match.group(0) if ans_match else None
    
    # 텍스트 길이를 기준으로 예전 번호(키)들을 정렬
    old_keys = list(old_opt_dict.keys())
    sorted_old_keys = sorted(old_keys, key=lambda k: len(old_opt_dict[k]))
    
    markers = ['①', '②', '③', '④', '⑤']
    old_to_new = {old_k: markers[i] for i, old_k in enumerate(sorted_old_keys)}
    
    # 1. 보기 텍스트 재구성
    new_opts_str = "\n"
    for i, old_k in enumerate(sorted_old_keys):
        new_opts_str += f"{markers[i]} {old_opt_dict[old_k]}\n"
        
    # 2. 정답 번호 재구성
    new_ans = old_to_new.get(old_ans, ans_text) if old_ans else ans_text
    
    # 3. 해설 내 번호 일괄 치환 (충돌 방지용 임시 태그 사용)
    new_exp_text = exp_text
    for old_k in old_keys:
        new_exp_text = new_exp_text.replace(old_k, f"__TEMP_{old_k}__")
    for old_k, new_k in old_to_new.items():
        new_exp_text = new_exp_text.replace(f"__TEMP_{old_k}__", new_k)
        
    return passage + new_opts_str.strip(), new_ans, new_exp_text

# 🚀 [수정됨] 모델 변수를 동적으로 받고 구글 네이티브 에러를 해결한 버전
def validate_batch_json(full_text, point_text, client, is_google_native, target_model):
    val_prompt = f"""당신은 엄격한 영어 문항 검수 위원입니다.
생성된 [문제 세트]가 아래 [체크리스트]를 만족하는지 검증하세요.

[체크리스트]
1. 난이도 조건: [META: ...] 태그의 난이도 배정 조건이 실제 설계에 반영되었는가?
2. 정답 무결성: ①보기 개수와 정답 번호가 일치하는지, ②정답 번호와 해설에서 설명하는 내용이 정확히 일치하는지, ③실제 선지 내용과 해설의 내용이 일치하는지? ④오답 검증 시, 반드시 [역할 B]에 제공된 '핵심 출제 포인트'를 다시 읽고(Reference), 해당 오답이 DB에 명시된 '대체 가능한 형태'나 '예외 규칙'에 해당하지 않는지 검증.
3. 물리적 모순 여부: '개수'를 묻는 문제인 경우, 선지(①~⑤)에 적힌 숫자의 최댓값이 [보기]에 제시된 문장(또는 단어)의 총개수를 초과하는지? (예: 문장은 3개인데 선지에 '4개', '5개'가 존재하면 즉시 F 처리할 것)
4. 맥락적 단서의 완전성: 시제, 수일치, 의미를 묻는 빈칸의 경우, 문장 내에 정답을 하나로 확정 지을 수 있는 명확한 단서(예: 시간 부사 등)가 존재하는지 비판적으로 검토할 것. 단서가 부족하여 다른 선지도 해석상 정답이 될 수 있는 논리적 틈이 있다면 즉시 F 처리할 것.

[문제 세트]
{full_text}

[출력 규칙 - 절대 엄수]
- 오직 유효한 JSON 형식으로만 출력. 마크다운이나 다른 설명 금지.
- 통과 = "P", 실패 = "F: [사유]"
- 예: {{"1": "P", "2": "F: 보기 시각적 균형 위반"}}
"""
    try:
        if is_google_native:
            import google.generativeai as genai
            clean_model = target_model.replace("google/", "") if "google/" in target_model else target_model
            model = genai.GenerativeModel(clean_model)
            response = model.generate_content(val_prompt)
            raw = response.text.strip()
        else:
            # 🚀 OpenRouter 400 에러의 주범인 'response_format' 강제 파라미터 삭제!
            response = client.chat.completions.create(
                model=target_model, 
                messages=[{"role": "user", "content": val_prompt}],
                temperature=0.0
            )
            raw = response.choices[0].message.content.strip()
            
        import re
        import json
        json_match = re.search(r'\{.*\}', raw, re.DOTALL)
        
        if json_match:
            try:
                return json.loads(json_match.group(0))
            except json.JSONDecodeError as parse_e:
                st.error(f"⚠️ JSON 파싱 실패: {parse_e}\nAI가 뱉은 원본 텍스트:\n{raw}")
                return {}
        else:
            st.error(f"⚠️ 검증기 JSON 추출 실패 (AI가 양식을 어겼습니다). 원본 응답:\n{raw}")
            return {}
            
    except Exception as e:
        # API 통신 자체가 터졌을 때 화면에 빨간 글씨로 띄워줌
        st.error(f"⚠️ 검증 함수 API 통신 에러 발생: {e}")
        return {}
        

# ── 데이터 로드 함수 (구글 시트 전용) ────────────────────────────────────────


@st.cache_data(ttl=180)
def load_gsheets_dual_db(q_url, c_url):
    try:
        def convert_to_csv_url(url):
            sheet_id_match = re.search(r'/d/([a-zA-Z0-9-_]+)', url)
            gid_match = re.search(r'gid=([0-9]+)', url)
            if sheet_id_match:
                sheet_id = sheet_id_match.group(1)
                gid = gid_match.group(1) if gid_match else "0"
                # 캐시 무력화 파라미터 포함
                return f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid={gid}&dummy={int(time.time())}"
            return url

        df_questions = pd.read_csv(convert_to_csv_url(q_url)).fillna('')
        df_concepts = pd.read_csv(convert_to_csv_url(c_url)).fillna('')
        
        # 열 이름 공백 제거
        df_questions.columns = [str(c).strip() for c in df_questions.columns]
        df_concepts.columns = [str(c).strip() for c in df_concepts.columns]
        # 🚀 SAFE DIAGNOSTIC PRINT 🚀
        st.session_state.db_columns = df_questions.columns.tolist()
        
        if '대분류' not in df_questions.columns:
            st.error("🚨 기출DB를 제대로 읽어오지 못했습니다. 구글 시트 공유 권한을 확인해 주세요.")
            return [], {}

        # 난이도 열 탐색
        diff_col = next((col for col in df_questions.columns if '난이도' in col), None)

        # 1. 'questions_db' 탭 파싱
        questions_pool = []
        for _, row in df_questions.iterrows():
            q_type = str(row.get('문제유형', '')).strip()
            if not q_type:
                continue
                
            q_diff = ""
            if diff_col:
                val = str(row.get(diff_col, '')).strip()
                if "상" in val: q_diff = "상"
                elif "중" in val: q_diff = "중"
                elif "하" in val: q_diff = "하"

            questions_pool.append({
                "u": str(row.get('대분류', '')).strip(),
                "s": str(row.get('소분류', '')).strip(),
                "t": q_type,
                "q": str(row.get('발문', '')).strip(),
                "c": str(row.get('보기', '')).strip(),
                "a": str(row.get('정답', '')).strip(),
                "e": str(row.get('해설', '')).strip(),
                "d": q_diff,
                "tag": str(row.get('태그', '')).strip() # 🚀 이 줄 추가
            })

        # 2. 'concept_hierarchy' 탭 파싱
        concepts_hierarchy = {}
        for _, row in df_concepts.iterrows():
            major = str(row.get('대분류', '')).strip()
            mid = str(row.get('중분류', '')).strip()
            minor = str(row.get('소분류', '')).strip()
            
            if not major or not mid:
                continue
                
            if major not in concepts_hierarchy:
                concepts_hierarchy[major] = {}
            if mid not in concepts_hierarchy[major]:
                concepts_hierarchy[major][mid] = []
                
            # 🚀 [수정됨] 프롬프트로 날아갈 통합 point_text를 여기서 미리 조립합니다 🚀
            # 선생님의 코드는 위에서 .fillna('')를 거치므로 빈 값은 빈 문자열('')로 처리됩니다.
            sub_category = str(row.get('소분류(카운팅 키)', minor)).strip()
            detail_rule = str(row.get('세부규칙(조회 키)', '')).strip()
            
            assembled_point = f"■ [{sub_category}] {detail_rule}\n"
            
            if str(row.get('출제포인트', '')).strip():
                assembled_point += f" - [출제포인트]: {str(row.get('출제포인트')).strip()}\n"
                
            if str(row.get('기본 규칙(정답 형태)', '')).strip():
                assembled_point += f" - [기본 규칙(정답 형태)]: {str(row.get('기본 규칙(정답 형태)')).strip()}\n"
                
            if str(row.get('☆(특수 or 예외 용법)', '')).strip():
                assembled_point += f" - [특수/예외 용법]: {str(row.get('☆(특수 or 예외 용법)')).strip()}\n"
                
            if str(row.get('인접 오류 형태', '')).strip():
                assembled_point += f" - [인접오류(매력적오답 타겟)]: {str(row.get('인접 오류 형태')).strip()}\n"
                
            if str(row.get('오류 설명', '')).strip():
                assembled_point += f" - [오류 사유/설명]: {str(row.get('오류 설명')).strip()}\n"
            
            # 조립이 완료된 거대한 텍스트를 "point" 키에 담아서 캐싱합니다.
            concepts_hierarchy[major][mid].append({
                "minor": minor,
                "difficulty": str(row.get('난이도', '')).strip(),
                "point": assembled_point 
            })
            
        return questions_pool, concepts_hierarchy
    except Exception as e:
        st.error(f"🚨 데이터 로드 중 치명적 오류 발생: {e}")
        return [], {}

# ── 사이드바 ───────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ API 설정")
    raw_api_key = st.text_input(
        "🔑 통합 API Key 입력창", 
        type="password", 
        placeholder="OpenRouter, Google AI, Azure 키 입력"
    )
    
    selected_model = st.selectbox(
        "🤖 출제 인공지능 엔진 모델",
        options=["anthropic/claude-sonnet-4.6", "anthropic/claude-opus-4.8", "openai/gpt-5.5", "openai/gpt-5.1", "google/gemini-3.1-pro-preview", "google/gemini-2.5-flash-lite"]
    )
    val_selected_model = st.selectbox("🔎 검수용 AI 모델", ["google/gemini-3.1-pro-preview", "google/gemini-2.5-pro", "anthropic/claude-sonnet-4.6", "anthropic/claude-opus-4.8", "openai/gpt-5.5", "openai/gpt-5.1"], help="전체 문항을 1회 검증할 때 사용할 모델을 선택합니다.")
    
    detected_platform = "대기 중..."
    if raw_api_key:
        if raw_api_key.startswith("sk-or-"):
            detected_platform = "🔗 OpenRouter 허브 모드"
        elif raw_api_key.startswith("AIzaSy"):
            detected_platform = "♊ Google AI Studio 모드"
        elif len(raw_api_key) == 32 or "azure" in raw_api_key.lower():
            detected_platform = "☁️ Microsoft Azure 모드"
        elif raw_api_key.startswith("sk-"):
            detected_platform = "🟢 OpenAI 직결 모드"
            
    if raw_api_key:
        st.caption(f"**활성화된 연결:** `{detected_platform}`")
# [여기서부터 복사하세요]
        if "Azure" in detected_platform:
            azure_endpoint = st.text_input("Azure Endpoint URL", placeholder="https://YOUR_RESOURCE.openai.azure.com/")
            
# ── 사이드바: 레벨 선택 UI (레버 스위치 & 배경색 동적 변경) ──
with st.sidebar:
    st.markdown("---")
    st.markdown("### 🎚️ 타겟 레벨 선택")
    
    # 1. 라디오 버튼 대신 '좌우로 움직이는 레버(Toggle)' 사용
    # 스위치가 꺼져있으면 False(H레벨), 켜져있으면 True(E레벨)
    IS_E_LEVEL = st.toggle("🟡 H 레벨 ↔ 🔵 E 레벨", value=False, help="스위치를 켜면 E레벨로 전환됩니다.")



    # 2. 선택된 레벨에 따른 사이드바 배경색 지정 (연노랑 / 연파랑)
    sidebar_bg_color = "#eef2ff" if IS_E_LEVEL else "#fffdf0" # E레벨: 연파랑, H레벨: 연노랑
    
    # 3. CSS 강제 주입을 통해 사이드바 배경색 실시간 변경 (부드러운 전환 효과 포함)
    st.markdown(f"""
        <style>
            [data-testid="stSidebar"] {{
                background-color: {sidebar_bg_color} !important;
                transition: background-color 0.4s ease-in-out;
            }}
        </style>
    """, unsafe_allow_html=True)

    
    EASY_COMBS, MID_COMBS, HARD_COMBS = get_difficulty_combs(IS_E_LEVEL)
# ── DEBUG PANEL ──
    with st.expander("🛠️ Developer Debug Panel", expanded=False):
        # 1. Cache Clear Button (Fixes stuck sheet data instantly)
        if st.button("🗑️ Clear GSheets Cache", use_container_width=True):
            st.cache_data.clear()
            st.rerun()
            
        # 2. Live Variables
        st.markdown("**Live State:**")
        st.write(f"Mode: {'E-Level' if IS_E_LEVEL else 'H-Level'}")
        st.write(f"Model: {selected_model}")
        
        # 3. API Response Output
        st.markdown("**Last API Raw Output:**")
        if "raw_api_log" in st.session_state:
            st.json(st.session_state.raw_api_log)
        else:
            st.caption("Awaiting first generation...")
            
        # 4. Filter Diagnostic Output
        st.markdown("**Last Filter Check:**")
        if "filter_log" in st.session_state:
            st.code(st.session_state.filter_log)
        
        # 5. DB Column Output (Updated)
        st.markdown("**Parsed DB Keys:**")
        if 'QUESTIONS' in locals() and QUESTIONS:
            st.write(list(QUESTIONS[0].keys()))
        else:
            st.write("Awaiting data load...")
            
        # 6. Validator Standalone Test (단독 검수기)
        st.markdown("---")
        st.markdown("**🔎 검수 로직 단독 테스트**")
        st.caption("생성 비용을 아끼고 검수 프롬프트만 튜닝할 때 사용하세요.")
        test_q_text = st.text_area("검수할 문제 텍스트 (복붙)", height=150, placeholder="【문제 1】\n...")
        test_p_text = st.text_area("참조할 DB 규칙 (선택)", height=80, placeholder="[인접오류] ...")
        
        if st.button("단독 검수 실행 (Test Validator)", use_container_width=True):
            safe_api_key = raw_api_key.strip() if raw_api_key else ""
            if not safe_api_key:
                st.error("상단에 API 키를 먼저 입력하세요.")
            elif not test_q_text:
                st.warning("검수할 문제 텍스트를 입력하세요.")
            else:
                # 임시 클라이언트 셋업
                tmp_client = None
                tmp_is_google_native = False
                
                if safe_api_key.startswith("AIzaSy"):
                    genai.configure(api_key=safe_api_key)
                    tmp_is_google_native = True
                elif safe_api_key.startswith("sk-or-"):
                    tmp_client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=safe_api_key)
                else:
                    tmp_client = OpenAI(api_key=safe_api_key)
                    
                with st.spinner(f"{val_selected_model} 모델로 검수 중..."):
                    val_res = validate_batch_json(
                        full_text=test_q_text, 
                        point_text=test_p_text, 
                        client=tmp_client, 
                        is_google_native=tmp_is_google_native, 
                        target_model=val_selected_model
                    )
                    st.success("검수 완료!")
                    st.json(val_res) # 결과를 예쁜 JSON 형태로 출력
            
# ── 사전 DB 로드 (캐싱) ──
# (사이드바 블록이 끝난 후, 제일 먼저 데이터를 불러와야 합니다!)
H_QUESTIONS, H_CONCEPTS = load_gsheets_dual_db(QUESTIONS_SHEET_URL, CONCEPTS_SHEET_URL)
E_QUESTIONS, E_CONCEPTS = load_gsheets_dual_db(E_QUESTIONS_SHEET_URL, E_CONCEPTS_SHEET_URL)

# 🚀 [복구] 지워졌던 전역 변수 스위칭 로직 다시 추가
if IS_E_LEVEL:
    QUESTIONS = E_QUESTIONS
    CONCEPTS = E_CONCEPTS
    build_generation_prompt = build_generation_prompt_e
else:
    QUESTIONS = H_QUESTIONS
    CONCEPTS = H_CONCEPTS
    build_generation_prompt = build_prompt_h

# (이 아래로는 기존 코드 그대로 유지)
PRIMARY_TYPES = [
    "어법상 맞는 것", "어법상 옳은 것", "어법상 옳지 않은 것",
    "빈칸 채우기", "개수 고르기", "올바른 영작",
    "어법상 어색한 것", "어법상 바른 것", "바르게 짝지어진 것",
]
ALL_TYPES = sorted(set(q["t"] for q in QUESTIONS if q["t"]))
SORTED_TYPES = [t for t in PRIMARY_TYPES if t in ALL_TYPES] + \
               [t for t in ALL_TYPES if t not in PRIMARY_TYPES]


# ── CSS ──────────────────────────────────────────────────
st.markdown("""
<style>
  .main-header {
    background: linear-gradient(135deg, #1e3a5f 0%, #2d6a9f 100%);
    color: white; padding: 1.8rem 2.5rem; border-radius: 12px; margin-bottom: 1.5rem;
  }
  .main-header h1 { margin: 0; font-size: 1.8rem; }
  .main-header p  { margin: 0.3rem 0 0; opacity: 0.85; font-size: 0.95rem; }
  .section-card {
    background: #f8fafc; border: 1px solid #e2e8f0;
    border-radius: 10px; padding: 1.2rem 1.5rem; margin-bottom: 0.8rem;
  }
  .question-box {
    background: white; border-left: 4px solid #2d6a9f;
    border-radius: 8px; padding: 1.4rem; margin: 0.8rem 0;
    box-shadow: 0 2px 8px rgba(0,0,0,0.06);
  }
  .answer-box {
    background: #f0fdf4; border: 1px solid #86efac;
    border-radius: 8px; padding: 0.8rem 1.2rem; margin-top: 0.6rem;
  }
  .explanation-box {
    background: #fffbeb; border: 1px solid #fcd34d;
    border-radius: 8px; padding: 0.8rem 1.2rem; margin-top: 0.4rem;
  }
  .stat-box {
    background: white; border: 1px solid #e2e8f0;
    border-radius: 8px; padding: 0.8rem; text-align: center;
  }
  .stat-box .num   { font-size: 1.8rem; font-weight: 700; color: #2d6a9f; }
  .stat-box .label { font-size: 0.8rem; color: #64748b; }
  .diff-badge {
    display: inline-block; border-radius: 20px;
    padding: 0.15rem 0.7rem; font-size: 0.8rem; font-weight: 600; margin-left: 0.4rem;
  }
  .diff-하   { background:#dbeafe; color:#1e40af; }
  .diff-중   { background:#dcfce7; color:#166534; }
  .diff-중상 { background:#fef9c3; color:#854d0e; }
  .diff-상   { background:#fee2e2; color:#991b1b; }
  .batch-card {
    background: #f0f9ff; border: 1px solid #bae6fd;
    border-radius: 10px; padding: 1rem 1.5rem; margin-bottom: 0.6rem;
  }
  .stButton > button {
    background: linear-gradient(135deg, #1e3a5f, #2d6a9f);
    color: white; border: none; border-radius: 8px;
    padding: 0.55rem 1.5rem; font-size: 0.95rem; font-weight: 600;
  }
  .stButton > button:hover { opacity: 0.9; }
  }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="main-header">
  <h1>📝 영어 기출 문제 생성기</h1>
  <p>강남구 중학교 기출 154문제 데이터 기반 · AI 응용 문제 자동 생성 (클라우드/로컬 하이브리드)</p>
</div>
""", unsafe_allow_html=True)

# ── 사이드바 통계 ──
with st.sidebar:
    st.markdown("---")
    st.markdown("### 📊 현재 로드된 DB 현황")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown(f'<div class="stat-box"><div class="num">{len(QUESTIONS)}</div><div class="label">기출/참고 데이터</div></div>', unsafe_allow_html=True)
    with c2:
        st.markdown(f'<div class="stat-box"><div class="num">{len(CONCEPTS)}</div><div class="label">챕터</div></div>', unsafe_allow_html=True)
    st.markdown("")
    st.markdown("**📚 문법 챕터**")
    for major in CONCEPTS:
        total = sum(len(v) for v in CONCEPTS[major].values())
        st.markdown(f"- {major} ({total}개 소개념)")

if "history" not in st.session_state:
    st.session_state.history = []
if "pending" not in st.session_state:
    st.session_state.pending = []

tab1, tab2, tab3 = st.tabs(["🤖 AI 문제 생성", "📚 기출 문제 탐색", "📋 생성 기록"])

# ════════════════════════════════════════════════════════
# TAB 1 : AI 문제 생성
# ════════════════════════════════════════════════════════
with tab1:
    col_left, col_right = st.columns([1, 1], gap="large")

    with col_left:
        st.markdown("### 📖 문법 개념 선택")
        st.markdown('<div class="section-card">', unsafe_allow_html=True)

        major_list = list(CONCEPTS.keys()) if CONCEPTS else ["데이터 없음"]
        selected_major = st.selectbox("① 챕터", major_list, key="major")

        mid_list = list(CONCEPTS[selected_major].keys()) if CONCEPTS and selected_major in CONCEPTS else ["데이터 없음"]
        selected_mid = st.selectbox("② Cell", mid_list, key="mid")

        minor_items = CONCEPTS[selected_major][selected_mid] if CONCEPTS and selected_major in CONCEPTS and selected_mid in CONCEPTS[selected_major] else []
        
        minor_options = [item["minor"] for item in minor_items if item["minor"]] if minor_items else []
        
        if minor_options:
            selected_minors = st.multiselect("③ 소분류 (복수 선택 가능)", minor_options, default=minor_options, key="minor")
            
            if not selected_minors:
                selected_minor_label = "선택 안 됨"
                diff = ""
                point_text = ""
                st.warning("소분류를 1개 이상 선택하세요.")
            elif len(selected_minors) == 1:
                selected_minor_label = selected_minors[0]
                selected_item = next((x for x in minor_items if x["minor"] == selected_minor_label), None)
                diff = selected_item["difficulty"] if selected_item else ""
                diff_class = f"diff-{diff}" if diff else ""
                point_text = selected_item.get("point", "")
                
                st.markdown(f'④ 개념의 난이도: <span class="diff-badge {diff_class}">{diff if diff else "미분류"}</span>', unsafe_allow_html=True)
                if point_text:
                    with st.expander("💡 출제 포인트 보기"):
                        st.markdown(point_text)
            else:
                selected_minor_label = ", ".join(selected_minors)
                diff = "복합"
                selected_items = [x for x in minor_items if x["minor"] in selected_minors]
                point_text = "\n\n".join([f"[{x['minor']}]\n{x.get('point', '')}" for x in selected_items if x.get("point")])
                
                st.markdown(f'④ 개념의 난이도: <span class="diff-badge diff-중상">복합 출제</span>', unsafe_allow_html=True)
                with st.expander("💡 복합 출제 포인트 보기"):
                    st.markdown(point_text)
        else:
            selected_minor_label = ""
            selected_minors = []
            diff = ""
            point_text = ""

        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("### ✏️ 추가 요청사항")
        extra = st.text_area(
            "추가 요청",
            placeholder="예) 난이도 상으로 함정 선지 포함\n예) 스포츠 주제로, 강남구 기출 스타일",
            height=100,
            label_visibility="collapsed",
        )

    with col_right:
        st.markdown("### 📋 문제 유형 & 개수")
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown("**문제 유형 선택 (복수 가능)**")

        safe_default = ["어법상 맞는 것"] if "어법상 맞는 것" in SORTED_TYPES else (SORTED_TYPES[:1] if SORTED_TYPES else None)

        selected_types = st.multiselect(
            "유형",
            SORTED_TYPES,
            default=safe_default,
            label_visibility="collapsed",
        )

        # ────────────────────────────────────────────────────────
        st.markdown("---")
        st.markdown("**📊 난이도 배정**")
        
        st.caption("배정 방식 선택")
        t_col1, t_col2, t_col3 = st.columns([2.5, 1, 6.5])
        with t_col1: 
            st.markdown("<div style='text-align:right; margin-top:4px; font-weight:bold; color:#475569;'>🤖 자동</div>", unsafe_allow_html=True)
        with t_col2: 
            is_manual = st.toggle("모드", value=False, key="auto_mode", label_visibility="collapsed")
        with t_col3: 
            st.markdown("<div style='margin-top:4px; font-weight:bold; color:#475569;'>🛠️ 수동</div>", unsafe_allow_html=True)

        is_disabled = not is_manual

        st.caption("난이도 활성화")
        btn_col1, btn_col2, btn_col3 = st.columns(3)
        with btn_col1: st.toggle("🔴 상", value=True, key="btn_high", disabled=is_disabled)
        with btn_col2: st.toggle("🔵 중", value=True, key="btn_mid", disabled=is_disabled)
        with btn_col3: st.toggle("🟢 하", value=True, key="btn_low", disabled=is_disabled)
        
        st.caption("문항 수 할당")
        num_col1, num_col2, num_col3 = st.columns(3)
        with num_col1: val_high = st.number_input("상 (개)", min_value=0, max_value=20, value=3, key="num_high", disabled=is_disabled)
        with num_col2: val_mid = st.number_input("중 (개)", min_value=0, max_value=20, value=4, key="num_mid", disabled=is_disabled)
        with num_col3: val_low = st.number_input("하 (개)", min_value=0, max_value=20, value=3, key="num_low", disabled=is_disabled)

        if not is_manual:
            final_high, final_mid, final_low = 3, 4, 3
            st.info("🤖 **[자동 모드]** 기본값(상3, 중4, 하3)으로 배정됩니다.")
        else:
            final_high, final_mid, final_low = val_high, val_mid, val_low
            
        # 🚀 [수정] 변수명을 직관적으로 바꾸고 안내 문구를 '총합'의 의미로 변경
        total_num = final_high + final_mid + final_low
        
        if is_manual:
            st.success(f"✅ 선택한 유형에 **총 {total_num}개**의 문제가 골고루 배정됩니다.")
        # ────────────────────────────────────────────────────────

        st.markdown('</div>', unsafe_allow_html=True)

        ref_pool = [q for q in QUESTIONS if selected_major in q["u"] or selected_mid in q.get("s","")]
        if not ref_pool:
            ref_pool = QUESTIONS
        st.info(f"📎 참고 기출: {len(ref_pool)}문제 ('{selected_major}' 관련)")
            
# ── 생성 버튼 ─────────────────────────────────────────
    st.markdown("---")
    gen_col1, gen_col2, gen_col3 = st.columns([2.5, 1.5, 1])
    
    with gen_col1:
        generate_btn = st.button("🚀 문제 생성하기", use_container_width=True)
        
    with gen_col2:
        # 🚀 toggle 대신 3단계 selectbox 사용. 공간 절약을 위해 라벨(제목) 숨김 처리
        validator_status = st.selectbox(
            "검증기 상태", 
            ["🛡️ 검증기 선택 (필수)", "❌ 검증기 끄기", "✅ 검증기 켜기"],
            index=0,
            label_visibility="collapsed" 
        )
        # 선택값에 따라 뒤에서 쓸 use_validator 값을 True/False로 자동 세팅
        use_validator = True if "켜기" in validator_status else False
        
    with gen_col3:
        clear_btn = st.button("🗑️ 결과 초기화", use_container_width=True)

    if clear_btn:
        st.session_state.pending = []
        st.rerun()

# ── 생성 실행 ─────────────────────────────────────────
    if generate_btn:
        
        # 🚀 [추가] '선택 (필수)' 상태인 채로 버튼을 누르면 팝업 띄우고 즉시 정지
        if "선택" in validator_status:
            show_validator_warning()
            st.stop()
            
        total_num = final_high + final_mid + final_low
        safe_api_key = raw_api_key.strip() if raw_api_key else ""
        
        if total_num == 0:
            st.error("⚠️ 생성할 문제 개수가 0개입니다. 난이도별 문항 수를 1개 이상 배정해주세요.")
        elif not safe_api_key:
            st.error("⚠️ 사이드바에 통합 API Key를 입력해주세요.")
        elif not selected_types:
            st.error("⚠️ 문제 유형을 하나 이상 선택해주세요.")
        else:
            batch_results = []
            progress = st.progress(0, text="생성 준비 중...")
            total = len(selected_types)
            
            # 🚀 [추가] 실시간 검수 로그를 저장할 독립 바구니 생성
            validation_debug_logs = []
            
            client = None
            is_google_native = False
            
            if safe_api_key.startswith("sk-or-"):
                client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=safe_api_key)
            elif safe_api_key.startswith("AIzaSy"):
                genai.configure(api_key=safe_api_key)
                is_google_native = True
            elif len(safe_api_key) == 32 or "azure" in safe_api_key.lower():
                try:
                    client = AzureOpenAI(api_key=safe_api_key, api_version="2024-02-15-preview", azure_endpoint=azure_endpoint)
                except NameError:
                    st.error("Azure 연결을 위한 엔드포인트 URL을 입력해주세요.")
            else:
                client = OpenAI(api_key=safe_api_key)

            # 레벨별 난이도 배정 메커니즘
            EASY_POOL = EASY_COMBS
            MID_POOL = MID_COMBS
            HARD_POOL = HARD_COMBS
            
            diff_targets = []
            for _ in range(final_high): diff_targets.append({"level": "상", "comb": random.choice(HARD_POOL)})
            for _ in range(final_mid):  diff_targets.append({"level": "중", "comb": random.choice(MID_POOL)})
            for _ in range(final_low):  diff_targets.append({"level": "하", "comb": random.choice(EASY_POOL)})
            
            allocations = {t: [] for t in selected_types}
            for i, diff_dict in enumerate(diff_targets):
                allocations[selected_types[i % len(selected_types)]].append(diff_dict)

            random.shuffle(TOPIC_LIST)
            topic_index = 0
            
            global_q_num = 1 
            all_generated_dict = {}
            all_qtype_map = {} 

            # ── 1. 통합 생성을 위한 1차 루프 ──
            for idx, qtype in enumerate(selected_types):
                type_diffs = allocations[qtype]
                if not type_diffs: continue 
                
                num_for_this_type = len(type_diffs)
                progress.progress((idx) / total, text=f"[{idx+1}/{total}] '{qtype}' 유형 {num_for_this_type}문제 생성 중...")

                type_matched = [q for q in QUESTIONS if q["t"] == qtype and "지문형" not in q.get("tag", "")]
                st.session_state.filter_log = f"Pool size for '{qtype}': {len(type_matched)} surviving questions."
                
                unit_matched = [q for q in QUESTIONS if selected_major in q["u"] and "지문형" not in q.get("tag", "")]
                qtype_pool = type_matched if len(type_matched) >= 3 else (unit_matched if unit_matched else [q for q in QUESTIONS if "지문형" not in q.get("tag", "")])
                ref_samples = random.sample(qtype_pool, min(4, len(qtype_pool)))
                ref_text = "\n\n".join([f"[기출 {i+1}]\n문제유형: {q['t']}\n발문: {q['q']}\n보기/지문: {q['c']}\n정답: {q['a']}" for i, q in enumerate(ref_samples)])

                integration_rule = ""
                if len(selected_minors) > 1:
                    integration_rule = f"12. [복합 출제 지시 (필수)]: 이번 세트는 여러 소분류 개념이 합쳐진 복합 테스트입니다. [역할 B]에 제시된 출제 포인트들을 반드시 골고루 활용하여 절대 특정 개념에만 편중되지 않도록 창작하세요.\n"
                
                q_assignments = ""
                for d_dict in type_diffs:
                    lvl = d_dict["level"]
                    a, b, c = d_dict["comb"]
                    topic = TOPIC_LIST[topic_index % len(TOPIC_LIST)]
                    topic_index += 1                    
                    
                    q_assignments += f"【문제 {global_q_num}】 타겟 난이도: [{lvl}] (조건: A={a}점, B={b}점, C={c}점) | 강제 지문 소재: [{topic}]\n"
                    global_q_num += 1

                prompt = build_generation_prompt(
                    ref_text=ref_text, selected_major=selected_major, selected_mid=selected_mid,
                    selected_minor_label=selected_minor_label, point_text=point_text, qtype=qtype,
                    num_for_this_type=num_for_this_type, extra=extra, q_assignments=q_assignments,
                    integration_rule=integration_rule
                )

                try:
                    if is_google_native:
                        model = genai.GenerativeModel(
                            "gemini-3.1-pro-preview",
                            generation_config=genai.types.GenerationConfig(
                                thinking_config=genai.types.ThinkingConfig(thinking_level="medium")
                            )
                        )
                        response = model.generate_content(prompt)
                        st.session_state.raw_api_log = response.model_dump()
                        try:
                            result_text = response.candidates[0].content.parts[-1].text
                        except Exception:
                            result_text = response.text
                    else:
                        response = client.chat.completions.create(
                            model=selected_model, messages=[{"role": "user", "content": prompt}],
                            temperature=0.75, max_tokens=9000
                        )
                        result_text = response.choices[0].message.content

                    if result_text:
                        result_text = result_text.replace(r"\_", "_").replace("&nbsp;", " ")
                    if result_text is None:
                        raise Exception("AI가 텍스트 대신 빈 값을 반환했습니다.")

                    raw_splits = re.split(r'(【문제 \d+】)', result_text)
                    for i in range(1, len(raw_splits), 2):
                        q_num = int(re.search(r'\d+', raw_splits[i]).group())
                        all_generated_dict[q_num] = raw_splits[i] + "\n" + raw_splits[i+1].strip()
                        all_qtype_map[q_num] = qtype
                
                except Exception as e:
                    st.error(f"{qtype} 유형 생성 중 에러 발생: {e}")

            # ── 2. 통합 검증 및 부분 재생성 루프 (1회 통신으로 모든 에러 교정) ──
            if use_validator and all_generated_dict:
                progress.progress(1.0, text="🔎 전체 문항 대상 통합 검수 진행 중...")
                combined_text = "\n\n".join([all_generated_dict[k] for k in sorted(all_generated_dict.keys())])
                
                # 단 1회의 검증 API 호출
                val_results = validate_batch_json(combined_text, point_text, client, is_google_native, val_selected_model)
                validation_debug_logs.append(f"📋 [시스템] 검수 AI가 보낸 원본 결과표: {val_results}")
                
                # 🚀 [핵심 변경] 반려된 문항들을 한 바구니에 수집
                failed_items = []
                for q_num_str, status in val_results.items():
                    status_str = str(status).strip()
                    
                    if status_str.startswith("F") or "실패" in status_str:
                        try:
                            match = re.search(r'\d+', str(q_num_str))
                            if not match: continue
                            q_num = int(match.group())
                            
                            fail_reason = status_str.replace("F:", "").replace("실패:", "").strip()
                            original_text = all_generated_dict.get(q_num, "")
                            
                            if original_text:
                                failed_items.append({
                                    "q_num": q_num,
                                    "fail_reason": fail_reason,
                                    "original_text": original_text
                                })
                        except Exception:
                            pass
                    else:
                        validation_debug_logs.append(f"✅ [통과] 문제 {q_num_str}번 문항은 무결성 검수 결과 이상이 없습니다.")
                
                # 🚀 바구니에 담긴 실패 문항이 있다면 단 1회의 통합 재생성 콜(Call 3) 전송
                if failed_items:
                    q_nums_str = ", ".join([str(item["q_num"]) for item in failed_items])
                    validation_debug_logs.append(f"🔄 [통합 교정 시작] 문제 {q_nums_str}번 반려 확인 -> 1회 통합 재생성 요청")
                    progress.progress(1.0, text=f"⚠️ 문제 {q_nums_str} 통합 재생성 중...")
                    
                    # 수집된 대상을 기반으로 묶음 프롬프트 조립
                    retry_prompt = build_retry_prompt(failed_items, point_text)
                    
                    try:
                        if is_google_native:
                            res = model.generate_content(retry_prompt)
                            try:
                                retry_result_text = res.candidates[0].content.parts[-1].text.strip()
                            except Exception:
                                retry_result_text = res.text.strip()
                        else:
                            response = client.chat.completions.create(
                                model=selected_model, 
                                messages=[{"role": "user", "content": retry_prompt}],
                                temperature=0.5, # 오답 디테일을 정교하게 고치기 위해 온도를 살짝 낮춤
                                max_tokens=9000
                            )
                            retry_result_text = response.choices[0].message.content.strip()
                        
                        # 🚀 돌아온 통합 수정본 텍스트를 문제 번호 태그 기준으로 다시 분할하여 원본 바구니 덮어쓰기
                        retry_splits = re.split(r'(【문제 \d+】)', retry_result_text)
                        for i in range(1, len(retry_splits), 2):
                            num_match = re.search(r'\d+', retry_splits[i])
                            if not num_match: continue
                            rq_num = int(num_match.group())
                            
                            # 매칭되는 원본의 실패 사유 추출
                            f_reason = next((item["fail_reason"] for item in failed_items if item["q_num"] == rq_num), "검수 기준 미달")
                            fixed_content = retry_splits[i] + "\n" + retry_splits[i+1].strip()
                            
                            # 최종 결과 바구니 갱신
                            all_generated_dict[rq_num] = f"🚨 **[육안 검수 요망: 재생성 문항 (사유: {f_reason})]**\n\n" + fixed_content
                            validation_debug_logs.append(f"✨ [교정 완료] 문제 {rq_num}번 문항의 통합 재작성이 완료되어 반영되었습니다.")
                            
                    except Exception as e:
                        validation_debug_logs.append(f"🔺 [통합 재생성 실패] API 통신 또는 파싱 오류 발생: {e}")
                        st.error(f"⚠️ 문제 {q_nums_str}번 통합 재생성 중 예외 발생: {e}")

            # ── 3. 최종 후처리 및 UI 파싱 준비 ── (이하 동일)

            # ── 3. 최종 후처리 및 UI 파싱 준비 ──
            progress.progress(1.0, text="✅ 최종 렌더링 준비 중...")
            current_idx = 1
            
            for q_num in sorted(all_generated_dict.keys()):
                full_text = all_generated_dict[q_num]
                qtype = all_qtype_map[q_num] 
                
                full_text = re.sub(r'\[META:.*?\]', '', full_text, flags=re.DOTALL)
                full_text = re.sub(r'【문제 \d+】', f'【문제 {current_idx}】', full_text)
                num_str = str(current_idx)
                current_idx += 1
                
                parts = {}
                tags = ["발문", "보기/지문", "정답", "해설"]
                for tag in tags:
                    key = f"[{tag}]"
                    if key in full_text:
                        start = full_text.index(key) + len(key)
                        nexts = [f"[{t}]" for t in tags if f"[{t}]" in full_text[start:]]
                        end = full_text.index(nexts[0], start) if nexts else len(full_text)
                        parts[tag] = full_text[start:end].replace("---", "").strip()

                if "보기/지문" in parts and "정답" in parts and "해설" in parts:
                    new_passage, new_ans, new_exp = sort_options(parts["보기/지문"], parts["정답"], parts["해설"])
                    dl_text = f"{num_str}. {parts.get('발문', '').strip()}\n\n{new_passage.strip()}\n\n정답: {new_ans.strip()}\n해설: {new_exp.strip()}\n"
                    full_text = f"【문제 {num_str}】\n[발문]\n{parts.get('발문', '')}\n\n[보기/지문]\n{new_passage}\n\n[정답]\n{new_ans}\n\n[해설]\n{new_exp}\n"
                else:
                    dl_text = full_text.replace("【", "").replace("】", ".").replace("[발문]\n", "").replace("[보기/지문]\n", "").replace("[정답]", "정답:").replace("[해설]", "해설:")

                is_valid = "🚨" not in full_text
                feedback = "검수 통과" if is_valid else "부분 재생성됨 (육안 확인 요망)"

                batch_results.append({
                    "type": qtype, "text": full_text, "dl_text": dl_text,
                    "is_valid": is_valid, "feedback": feedback
                })

            # ── 4. 세션 히스토리에 최종 저장 ──
            entry = {
                "major": selected_major,
                "mid": selected_mid,
                "minor": selected_minor_label,
                "difficulty": f"총 {total_num}문제 (상{final_high}/중{final_mid}/하{final_low})",
                "types": selected_types,
                "count": total_num,
                "results": batch_results,
                "val_logs": validation_debug_logs  # 🚀 수집한 실시간 로그를 보따리에 함께 패킹!
            }
            st.session_state.history.append(entry)
            st.session_state.pending = [entry]
            
            st.rerun()  # 🚀 안심하고 새로고침 실행 (세션에 들어가 있으므로 증발하지 않음)

    # ── 결과 표시 및 다운로드 UI ─────────────────────────────────────────
    if st.session_state.pending:
        entry = st.session_state.pending[-1]
        st.markdown("---")
        st.markdown(f"### 📄 생성 결과 — {entry['major']} > {entry['mid']} > {entry['minor']} ({entry.get('difficulty', '')})")

        # 🚀 [핵심 변경] 상단에 독립된 화면 공간을 뚫어 검수 로그 박스를 배치합니다.
        if "val_logs" in entry and entry["val_logs"]:
            with st.expander("🔍 AI 문항 실시간 검수 및 자동 교정 로그 (자세히 보기)", expanded=True):
                for log in entry["val_logs"]:
                    if "🔄" in log:
                        st.warning(log)
                    elif "❌" in log or "🔺" in log:
                        st.error(log)
                    elif "✨" in log or "✅" in log:
                        st.success(log)
                    else:
                        st.info(log)

        prev_type = None
        for res in entry["results"]:
            if res['type'] != prev_type:
                st.markdown(f"### 🟦 【{res['type']}】 유형")
                prev_type = res['type']
            
            if not res.get("is_valid", True):
                st.error(f"⚠️ **{res.get('feedback')}**")
            else:
                st.success("✅ **검수 통과**")

            raw = str(res.get("text", ""))
            
            if raw.startswith("[통신오류]"):
                st.error("⚠️ 통신 에러가 발생하여 생성이 중단되었습니다.")
                st.warning(raw)
                continue 
                
            problems = raw.split("【문제")
            if len(problems) <= 1:
                st.warning("⚠️ 양식이 깨졌거나 렌더링 오류가 발생했습니다. 원본을 확인하세요.")
                st.markdown(raw.replace("\n", "  \n"))
                continue

            valid_problem_found = False
            for prob in problems:
                prob = prob.strip()
                if not prob or not prob[0].isdigit():
                    continue
                
                valid_problem_found = True
                full = "【문제" + prob
                parts = {}
                tags = ["발문", "보기/지문", "정답", "해설"]
                for tag in tags:
                    key = f"[{tag}]"
                    if key in full:
                        start = full.index(key) + len(key)
                        nexts = [f"[{t}]" for t in tags if f"[{t}]" in full[start:]]
                        end = full.index(nexts[0], start) if nexts else len(full)
                        parts[tag] = full[start:end].replace("---", "").strip()

                st.markdown('<div class="question-box">', unsafe_allow_html=True)
                num_part = prob[:10].split("】")[0]
                st.markdown(f"**【문제{num_part}】**")
                if "발문" in parts:
                    st.markdown(f"**{parts['발문']}**", unsafe_allow_html=True)
                if "보기/지문" in parts:
                    st.markdown(parts["보기/지문"], unsafe_allow_html=True)
                if "정답" in parts:
                    st.markdown(f'<div class="answer-box">✅ <b>정답:</b> {parts["정답"]}</div>', unsafe_allow_html=True)
                if "해설" in parts:
                    st.markdown("💡 **해설:**")
                    st.markdown(parts["해설"].replace("\n", "  \n"))
                st.markdown('</div>', unsafe_allow_html=True)
            
            if not valid_problem_found:
                st.markdown(raw.replace("\n", "  \n"))

        st.markdown("---")
        set_text = f"[{entry['major']} > {entry['mid']} > {entry['minor']}] {entry.get('difficulty', '')}\n\n"
        
        prev_dl_type = None
        for r in entry["results"]:
            if r['type'] != prev_dl_type:
                set_text += f"\n🟦 {r['type']} 유형\n\n"
                prev_dl_type = r['type']
            set_text += f"{r.get('dl_text', r['text'])}\n\n"
            
        now_str = datetime.now().strftime("%y%m%d")
        safe_model = selected_model.split('/')[-1] 
        f_name = f"{entry['major']}_{entry['mid']}_{now_str}_{safe_model}"
    
        dl_col1, dl_col2 = st.columns(2)
        unique_key = len(st.session_state.history) 
        
        with dl_col1:
            st.download_button(
                "⬇️ 방금 만든 문제 다운로드 (.txt)",
                data=set_text.encode("utf-8"), file_name=f"{f_name}.txt",
                mime="text/plain", use_container_width=True, key=f"dl_current_txt_{unique_key}" 
            )
        with dl_col2:
            st.download_button(
                "📄 방금 만든 문제 다운로드 (.docx)",
                data=create_word_document(entry, is_multiple=False), file_name=f"{f_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, key=f"dl_current_docx_{unique_key}"
            )
# ════════════════════════════════════════════════════════
# TAB 2 : 기출 문제 탐색
# ════════════════════════════════════════════════════════

with tab2:
    st.markdown("### 📚 기출 문제 탐색")

    fc1, fc2, fc3 = st.columns(3)
    with fc1:
        f_major = st.selectbox("챕터 필터", ["전체"] + list(CONCEPTS.keys()), key="f_major")
    with fc2:
        all_q_types = sorted(set(q["t"] for q in QUESTIONS if q["t"]))
        f_type = st.selectbox("문제유형 필터", ["전체"] + all_q_types, key="f_type")
    with fc3:
        kw = st.text_input("키워드", placeholder="예) 관계대명사, 현재완료")

    filtered = QUESTIONS
    if f_major != "전체":
        filtered = [q for q in filtered if f_major in q["u"]]
    if f_type != "전체":
        filtered = [q for q in filtered if q["t"] == f_type]
    if kw:
        kl = kw.lower()
        filtered = [q for q in filtered if kl in q["q"].lower() or kl in q["s"].lower() or kl in q["c"].lower()]

    def render_newlines(text: str) -> str:
        return text.replace("\n", "  \n")

    st.markdown(f"**검색 결과: {len(filtered)}문제**")
    for q in filtered[:30]:
        with st.expander(f"[{q['u']} > {q['s']}] {q['t']} — {q['q'][:60]}..."):
            st.markdown(f"**발문:** {q['q']}")
            if q["c"]:
                st.markdown("**보기/지문:**")
                st.markdown(render_newlines(q["c"]))
            st.markdown(f'<div class="answer-box">✅ <b>정답:</b> {q["a"]}</div>', unsafe_allow_html=True)
            if q["e"]:
                st.markdown("💡 **해설:**")
                st.markdown(render_newlines(q["e"]))
    if len(filtered) > 30:
        st.info("상위 30개만 표시됩니다. 필터를 좁혀 검색하세요.")


# ════════════════════════════════════════════════════════
# TAB 3 : 생성 기록
# ════════════════════════════════════════════════════════
with tab3:
    st.markdown("### 📋 생성 기록")

    if not st.session_state.history:
        st.info("아직 생성된 문제가 없습니다.")
    else:
        # 🚀 [수정 1] 전체 통합 다운로드 (txt / docx 분리)
        all_combined = ""
        for i, h in enumerate(st.session_state.history):
            all_combined += f"\n{'='*70}\n"
            all_combined += f"[세트 {i+1}] {h['major']} > {h['mid']} > {h['minor']} | 난이도: {h['difficulty']}\n"
            all_combined += f"{'='*70}\n\n"
            
            prev_dl_type = None
            for r in h["results"]:
                # 기존: all_combined += f"【{r['type']} 유형】\n\n{r['text']}\n\n"
                
                # 🚀 이렇게 교체
                if r['type'] != prev_dl_type:
                    all_combined += f"\n🟦 {r['type']} 유형\n\n"
                    prev_dl_type = r['type']
                all_combined += f"{r.get('dl_text', r['text'])}\n\n"

        c1, c2 = st.columns(2)
        with c1:
            st.download_button(
                "⬇️ 전체 기록 통합 다운로드 (.txt)",
                data=all_combined.encode("utf-8"),
                file_name="전체_생성문제.txt",
                mime="text/plain",
                use_container_width=True,
                key="dl_history_all_txt"
            )
        with c2:
            st.download_button(
                "📄 전체 기록 통합 다운로드 (.docx)",
                data=create_word_document(st.session_state.history, is_multiple=True),
                file_name="전체_생성문제.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_history_all_docx"
            )
        st.markdown("---")

        for i, h in enumerate(reversed(st.session_state.history)):
            idx = len(st.session_state.history) - i
            label = f"세트 {idx} | {h['major']} > {h['minor']} | {', '.join(h['types'])} | 각 {h['count']}문제"
            with st.expander(label):
                for r in h["results"]:
                    st.markdown(f"**▶ {r['type']} 유형**")
                    st.markdown(r["text"])
                    st.markdown("---")

                # 기존 파일명 변수 교체
                set_f_name = f"{datetime.now().strftime('%y%m%d')}_{selected_model.split('/')[-1]}_{h['major']}_{h['mid']}"
                
                sc1, sc2 = st.columns(2)
                with sc1:
                    st.download_button(
                        f"⬇️ 세트 {idx} 다운로드 (.txt)",
                        data=set_text.encode("utf-8"),
                        file_name=f"{set_f_name}.txt",
                        mime="text/plain",
                        use_container_width=True,
                        key=f"dl_history_set_txt_{idx}",
                    )
                with sc2:
                    st.download_button(
                        f"📄 세트 {idx} 다운로드 (.docx)",
                        data=create_word_document(h, is_multiple=False),
                        file_name=f"{set_f_name}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key=f"dl_history_set_docx_{idx}",
                    )
